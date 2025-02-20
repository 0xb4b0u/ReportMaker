import os

import docx

from api_requests import (
    parse_vuln_info,
    get_vuln_info,
    get_description,
    get_potential_false_positive,
    remove_false_positive,
)


# Récupérer les informations de la CVE
CVE_NAME = str(input("Entrez le nom de la CVE: "))

affected_devices = parse_vuln_info(get_vuln_info(CVE_NAME))
false_positive = get_potential_false_positive(affected_devices)
affected_devices = remove_false_positive(affected_devices, false_positive)
description = get_description(CVE_NAME)

# Création du document Word
report = docx.Document()

# Ajour du titre du rapport
report.add_heading(f"Rapport pour la {CVE_NAME}", 0)

# Présenter les informations des machines impactées
report.add_heading(f"{len(affected_devices)} machine(s) impactée(s)", level=1)
report.add_paragraph("Tableau des machines impactées par la CVE")

# Création du tableau
table = report.add_table(rows=1, cols=5)
table.style = "Table Grid"

# Ajout des en-têtes
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Instance"
hdr_cells[1].text = "IP"
hdr_cells[2].text = "Profile"
hdr_cells[3].text = "Site"
hdr_cells[4].text = "Dernière détection"

# Ajout des données au tableau
for key, value in affected_devices.items():
    row_cells = table.add_row().cells
    row_cells[0].text = key
    row_cells[1].text = value[0][0]
    row_cells[2].text = value[0][1]
    row_cells[3].text = value[0][2]
    row_cells[4].text = value[0][3]

# Présenter les informations des machines impactées
report.add_heading(
    f"{len(false_positive)} machines considérées comme faux positif", level=1
)
report.add_paragraph("Tableau des machines considérées comme faux positif")

# Création du tableau
table2 = report.add_table(rows=1, cols=5)
table2.style = "Table Grid"

# Ajout des en-têtes
hdr_cells = table2.rows[0].cells
hdr_cells[0].text = "Instance"
hdr_cells[1].text = "IP"
hdr_cells[2].text = "Profile"
hdr_cells[3].text = "Site"
hdr_cells[4].text = "Dernière détection"

# Ajout des données au tableau
for key, value in false_positive.items():
    row_cells = table2.add_row().cells
    row_cells[0].text = key
    row_cells[1].text = value[0]
    row_cells[2].text = value[1]
    row_cells[3].text = value[2]
    row_cells[4].text = value[3]


# Ajout de la description de la CVE
report.add_heading("Description de la CVE", level=1)
report.add_paragraph(description)

report.add_heading("Solution(s) à mettre en place", level=1)

# Creation du repertoire reports
reports_dir = "reports"
os.makedirs(reports_dir, exist_ok=True)

# Sauvegarde du document Word
docx_path = os.path.join(reports_dir, f"{CVE_NAME}_report.docx")
report.save(docx_path)
